"""Core document generation service for fee proposals."""

import io
import os
import re
from datetime import date, datetime, timedelta
import json

from mailmerge import MailMerge
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm

from services.fee_calculator import (
    number_to_word, get_ordinal_suffix, get_legislation,
    format_riba_stages, determine_riba_stages, build_input_data, get_initials
)
from services import fee_text_templates as txt

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "styles.docx")
SIGNATURES_DIR = os.path.join(BASE_DIR, "signatures")


def _add_blank_line(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph("")
        p.style = doc.styles["Standard_Text_No_Line_Spacing"]


def _add_para(doc, text, style):
    p = doc.add_paragraph(text)
    p.style = doc.styles[style]
    return p


def _add_meeting_text(doc, service_data, prefix="Where necessary, attend meetings with the design team and approving authorities to present and discuss the outline strategy proposals.", unlimited="Attend all relevant meetings with the design team and approving authorities to present and discuss the outline strategy proposals."):
    if service_data.get("meeting_limit"):
        meetlim = service_data.get("meeting_number")
        if meetlim and int(meetlim) < 11:
            meet_num = number_to_word(int(meetlim))
        else:
            meet_num = str(meetlim) if meetlim else "N/A"
        _add_para(doc, f"{prefix} Allowance has been made in this proposal for attendance at {meet_num} project meetings.", "ASTUTE_Bullet_Points")
    else:
        _add_para(doc, unlimited, "ASTUTE_Bullet_Points")


def _find_service(input_data, key):
    return next((s for s in input_data if s["key"] == key), None)


def _get_opt(svc_data):
    return " (Optional)" if svc_data and svc_data.get("optional") else ""


def _get_num_models_text(svc_data, default="three"):
    nm = svc_data.get("num_models") if svc_data else None
    if nm and int(nm) <= 10:
        return number_to_word(int(nm))
    return str(nm) if nm else default


def generate_proposal(data) -> io.BytesIO:
    first_name = data.client.first_name
    surname = data.client.surname
    address_lines = data.client.address_lines
    project_name = data.project.project_name
    project_location = data.project.project_location
    project_country = data.project.country.value
    provide_hourly_rates = data.fee_options.include_hourly_rates
    PII = data.fee_options.pii_limit
    engineer = data.fee_options.engineer_name
    s14 = data.design_stages_1_4
    s5 = data.design_stages_5
    s6 = data.design_stages_6

    engineers_path = os.path.join(BASE_DIR, "data", "engineers.json")
    with open(engineers_path, "r", encoding="utf-8") as f:
        engineers = json.load(f)
    engineer_row = next((e for e in engineers if e["full_name"] == engineer), None)
    if not engineer_row:
        raise ValueError(f"Engineer '{engineer}' not found")

    job_title = engineer_row["job_title"]
    telno = engineer_row["phone_number"]
    emailad = f"{engineer_row['email_prefix']}@firedynamicsgroup.com"
    pic_name = re.sub(r'\W+', '', engineer).lower() + ".jpg"
    sig_path = os.path.join(SIGNATURES_DIR, pic_name)
    has_signature = os.path.exists(sig_path)
    initials = get_initials(engineer)
    legislation = get_legislation(project_country)
    input_data = build_input_data(data)

    address = txt.OFFICE_ADDRESS
    temp_output = io.BytesIO()
    with MailMerge(TEMPLATE_PATH) as mm_doc:
        mm_doc.merge(add1=address[0], add2=address[1], add3=address[2], add4=address[3], email=emailad, tel=telno)
        mm_doc.write(temp_output)
    temp_output.seek(0)
    doc = Document(temp_output)

    # DATE
    _add_blank_line(doc, 1)
    today = date.today()
    month = today.strftime("%B")
    year = today.strftime("%Y")
    day = today.day
    p = _add_para(doc, str(day), "Standard_Text_No_Line_Spacing")
    run = p.add_run(get_ordinal_suffix(day))
    run.font.superscript = True
    p.add_run(f" {month}, {year}")

    # CLIENT ADDRESS
    _add_blank_line(doc, 2)
    _add_para(doc, f"{first_name} {surname}", "Standard_Text")
    for line in address_lines:
        if line.strip():
            _add_para(doc, line, "Standard_Text_No_Line_Spacing")
    _add_blank_line(doc, 3)

    # INTRODUCTION
    _add_para(doc, f"Dear {first_name},", "Standard_Text")
    _add_para(doc, f"REF: {project_name} - Fire Engineering Fee Proposal", "Ref:")
    p = _add_para(doc, "Further to recent discussions, please find detailed below our fees for providing specialist fire safety engineering advice", "Standard_Text")
    if project_location in project_name and "Project" not in project_name:
        p.add_run(f" for the proposed {project_name} project.")
    elif project_location in project_name and "Project" in project_name:
        p.add_run(f" for the proposed {project_name}.")
    elif "Project" in project_name:
        p.add_run(f" for the proposed {project_name} in {project_location}.")
    else:
        p.add_run(f" for the proposed works at {project_name} in {project_location}.")

    third_party = s14.peer_review.included
    if not third_party:
        riba_stages = determine_riba_stages(data)
        riba_text = format_riba_stages(riba_stages)
        p.add_run(f"It is understood that our involvement is required during RIBA {riba_text}.")
    else:
        nm = s14.peer_review.num_models
        nm_text = number_to_word(nm) if nm and nm <= 10 else (str(nm) if nm else "N/A")
        _add_para(doc, f"We understand that the extent of our involvement will be to provide an independent third party review of a CFD Modelling analysis undertaken as part of the fire engineering design of the project. This will require the review of the input and output data of {nm_text} CFD models as well as a review of the accompanying report.", "Standard_Text")

    if s14.common_corridor_cfd.included:
        if not s14.common_corridor_cfd.extended_travel_distance:
            _add_para(doc, txt.INTRO_COMMON_CORRIDOR_DEPRESSURISATION, "Standard_Text")
        else:
            _add_para(doc, txt.INTRO_COMMON_CORRIDOR_EXTENDED_TRAVEL, "Standard_Text")
    if s14.open_plan_cfd.included:
        _add_para(doc, txt.INTRO_OPEN_PLAN, "Standard_Text")
    if s14.warehouse_cfd.included:
        _add_para(doc, txt.INTRO_WAREHOUSE_CFD, "Standard_Text")
    if s14.warehouse_structural.included:
        _add_para(doc, txt.INTRO_WAREHOUSE_STRUCTURAL, "Standard_Text")
    if s14.london_plan.included and s14.gateway.included:
        _add_para(doc, txt.INTRO_LONDON_PLAN_AND_GATEWAY, "Standard_Text")
    elif s14.london_plan.included:
        _add_para(doc, txt.INTRO_LONDON_PLAN_ONLY, "Standard_Text")
    elif s14.gateway.included:
        _add_para(doc, txt.INTRO_GATEWAY_ONLY, "Standard_Text")
    _add_para(doc, txt.INTRO_APPENDIX_REFERENCE, "Standard_Text")

    # FEES
    _add_para(doc, "Fees", "Subheading")
    if len(input_data) == 1:
        item = input_data[0]
        p = _add_para(doc, "Our proposed fee for the scope outlined in this document is ", "Standard_Text")
        run = p.add_run("\u00a3{:0,.2f}".format(float(item["fee"])))
        run.font.bold = True
        p.add_run("." if project_country == "J" else " exc. VAT.")
    else:
        _add_para(doc, "Our proposed fees for the scope outlined in this document are as follows:", "Standard_Text")
        fee_total = 0
        table = doc.add_table(0, 0)
        table.style = "invisible"
        table.add_column(Cm(12))
        table.add_column(Cm(5))
        for n, item in enumerate(input_data):
            opt = " (Optional)" if item["optional"] else ""
            end = f" (Up to {item['end_date']})" if item["end_date"] else ""
            table.add_row()
            row = table.rows[n]
            cell_a = row.cells[0]
            p_a = cell_a.paragraphs[0]
            p_a.text = f"{item['ref']}{opt}{end}"
            p_a.style = doc.styles["Standard_Text"]
            cell_b = row.cells[1]
            p_b = cell_b.paragraphs[0]
            p_b.text = "\u00a3{:0,.2f}".format(float(item["fee"]))
            p_b.style = doc.styles["Standard_Text"]
            fee_total += float(item["fee"])
        n = len(input_data)
        table.add_row()
        row = table.rows[n]
        row.cells[0].text = "Total"
        row.cells[1].text = "\u00a3{:0,.2f}".format(fee_total)
        table.rows[n].cells[0].paragraphs[0].style = doc.styles["Standard_Text"]
        table.rows[n].cells[1].paragraphs[0].style = doc.styles["Standard_Text"]
        table.rows[n].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[n].cells[1].paragraphs[0].runs[0].font.bold = True
        if project_country != "J":
            _add_para(doc, "All fees quoted are exclusive of VAT. ", "Standard_Text")

    if not provide_hourly_rates:
        _add_para(doc, "Any additional works beyond the detailed scope will either be charged at our standard hourly rates (available upon request) or would be subject to an additional fee agreement.", "Standard_Text")
    else:
        p = _add_para(doc, "Any additional works beyond the detailed scope will either be charged at our hourly rates, ", "Standard_Text")
        if project_country != "J":
            p.add_run("exclusive of VAT, ")
        p.add_run("as follows:")
        table = doc.add_table(0, 0)
        table.style = "invisible"
        table.add_column(Cm(8))
        table.add_column(Cm(5))
        for n, (title, rate) in enumerate(txt.HOURLY_RATES):
            table.add_row()
            row = table.rows[n]
            row.cells[0].text = title
            row.cells[0].paragraphs[0].style = doc.styles["Standard_Text"]
            row.cells[1].text = rate
            row.cells[1].paragraphs[0].style = doc.styles["Standard_Text"]

    # TERMS
    validity_date = (datetime.now() + timedelta(days=60)).strftime("%d/%m/%y")
    PII_fmt = "{:0,.0f}".format(PII)
    vat_text = "" if project_country == "J" else "+VAT"
    _add_para(doc, "Terms of Business", "Subheading")
    _add_para(doc, f"We propose to offer professional indemnity insurance to the value of \u00a3{PII_fmt} in the aggregate for this project. If a higher level of PII is required, this will need to be agreed prior to commencement of the works.", "Standard_Text")
    _add_para(doc, txt.TERMS_ACE.format(vat_text=vat_text), "Standard_Text")
    _add_para(doc, txt.TERMS_INVOICING, "Standard_Text")
    _add_para(doc, f"This fee proposal is valid until {validity_date} (60 days).", "Standard_Text")

    # CLOSING
    _add_blank_line(doc, 2)
    _add_para(doc, txt.TERMS_CLOSING, "Standard_Text")
    _add_blank_line(doc, 2)
    _add_para(doc, "Yours Sincerely ", "Standard_Text_No_Line_Spacing")
    _add_blank_line(doc, 1)
    p = _add_para(doc, "", "Standard_Text_No_Line_Spacing")
    r = p.add_run()
    if has_signature:
        r.add_picture(sig_path, height=Cm(2.5))
    _add_blank_line(doc, 1)
    _add_para(doc, engineer, "Sig")
    _add_para(doc, job_title, "Sig")
    _add_para(doc, "Fire Dynamics", "Sig")

    # HEADER
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = f"Fire Engineering Fee Proposal \r{project_name}"

    # PAGE BREAK -> APPENDIX A
    p = _add_para(doc, "", "Standard_Text")
    p.add_run().add_break(WD_BREAK.PAGE)
    _add_para(doc, "Appendix A - Detailed Scope of Works", "Subheading")
    _write_appendix_a(doc, data, input_data, legislation)

    # PAGE BREAK -> APPENDIX B
    p = _add_para(doc, "", "Standard_Text")
    p.add_run().add_break(WD_BREAK.PAGE)
    _add_para(doc, "Appendix B - Exclusions", "Subheading")
    _write_appendix_b(doc, data, input_data)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _write_appendix_a(doc, data, input_data, legislation):
    s14 = data.design_stages_1_4
    s5 = data.design_stages_5
    s6 = data.design_stages_6
    stage_1 = _find_service(input_data, "stage_1")
    stage_2 = _find_service(input_data, "stage_2")
    lp = _find_service(input_data, "london_plan")
    gw = _find_service(input_data, "gateway")
    stage_3 = _find_service(input_data, "stage_3")
    stage_4 = _find_service(input_data, "stage_4")
    cc = _find_service(input_data, "common_corridor_cfd")
    op = _find_service(input_data, "open_plan_cfd")
    wcfd = _find_service(input_data, "warehouse_cfd")
    ws = _find_service(input_data, "warehouse_structural")
    pr = _find_service(input_data, "peer_review")
    ca = _find_service(input_data, "construction_advice")
    sv = _find_service(input_data, "site_visits")
    po = _find_service(input_data, "phased_occupation")
    cfsmp = _find_service(input_data, "cfsmp")
    sr = _find_service(input_data, "site_risk_assessment")
    r38 = _find_service(input_data, "regulation_38")
    ews1 = _find_service(input_data, "ews1_forms")
    rro = _find_service(input_data, "rro_risk_assessment")
    cm = _find_service(input_data, "client_monitoring")

    if stage_1:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 1{_get_opt(stage_1)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        for t in txt.STAGE_1_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.STAGE_1_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, stage_1, "Where necessary, attend meetings with the project team and relevant stakeholders. ", "Attend all relevant meetings with the project team and relevant stakeholders")

    if stage_2:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 2{_get_opt(stage_2)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        for t in txt.STAGE_2_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STAGE_2_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.STAGE_2_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, stage_2)

    if lp:
        _add_para(doc, f"Scope of Works \u2013 London Plan Fire Statement{_get_opt(lp)}", "ASTUTE SubHeader")
        for t in txt.LONDON_PLAN_INTRO:
            _add_para(doc, t, "Standard_Text")
        for t in txt.LONDON_PLAN_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.LONDON_PLAN_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.LONDON_PLAN_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, lp)

    if gw:
        _add_para(doc, f"Scope of Works \u2013 Gateway 1 Fire Statement{_get_opt(gw)}", "ASTUTE SubHeader")
        for t in txt.GATEWAY_INTRO:
            _add_para(doc, t, "Standard_Text")
        for t in txt.GATEWAY_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.GATEWAY_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.GATEWAY_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, gw)

    if stage_3:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 3{_get_opt(stage_3)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        if stage_2:
            for t in txt.STAGE_3_SCOPE_WITH_STAGE_2:
                _add_para(doc, t, "ASTUTE_Bullet_Points")
        else:
            for t in txt.STAGE_3_SCOPE_WITHOUT_STAGE_2:
                _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STAGE_3_COMMON_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STAGE_3_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        _add_para(doc, txt.STAGE_3_DELIVERABLES_TEMPLATE.format(legislation=legislation), "ASTUTE_Bullet_Points")
        for t in txt.STAGE_3_DELIVERABLES_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_meeting_text(doc, stage_3)

    if stage_4:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 4{_get_opt(stage_4)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        if not stage_3:
            for t in txt.STAGE_4_SCOPE_WITHOUT_STAGE_3:
                _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STAGE_4_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        _add_para(doc, txt.STAGE_4_DELIVERABLES_TEMPLATE.format(legislation=legislation), "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, stage_4)

    if cc:
        nmt = _get_num_models_text(cc)
        _add_para(doc, f"Scope of Works - CFD Modelling Study of Common Corridors{_get_opt(cc)}", "ASTUTE SubHeader")
        _add_para(doc, txt.COMMON_CORRIDOR_INTRO, "Standard_Text")
        _add_para(doc, txt.COMMON_CORRIDOR_SOLUTION, "Standard_Text")
        _add_para(doc, txt.COMMON_CORRIDOR_CFD_INTRO, "Standard_Text")
        for t in txt.COMMON_CORRIDOR_SCOPE:
            _add_para(doc, t.format(num_models=nmt), "ASTUTE_Bullet_Points")

    if op:
        nmt = _get_num_models_text(op)
        _add_para(doc, f"Scope of Works - CFD Modelling of Open Plan Apartments{_get_opt(op)}", "ASTUTE SubHeader")
        _add_para(doc, txt.OPEN_PLAN_INTRO_TEMPLATE.format(legislation=legislation), "Standard_Text")
        for t in txt.OPEN_PLAN_METHODOLOGY:
            _add_para(doc, t, "Standard_Text")
        _add_para(doc, txt.OPEN_PLAN_ANTICIPATED_TEMPLATE.format(legislation=legislation), "Standard_Text")
        _add_para(doc, "The scope of work for this study would be as follows: ", "Standard_Text")
        for t in txt.OPEN_PLAN_SCOPE:
            _add_para(doc, t.format(num_models=nmt), "ASTUTE_Bullet_Points")

    if wcfd:
        nmt = _get_num_models_text(wcfd)
        _add_para(doc, f"Scope of Works - Warehouse CFD Modelling Study{_get_opt(wcfd)}", "ASTUTE SubHeader")
        _add_para(doc, txt.WAREHOUSE_CFD_INTRO, "Standard_Text")
        for t in txt.WAREHOUSE_CFD_SCOPE:
            _add_para(doc, t.format(num_models=nmt), "ASTUTE_Bullet_Points")

    if ws:
        _add_para(doc, f"Scope of Works - Structural Fire Engineering Assessment{_get_opt(ws)}", "ASTUTE SubHeader")
        _add_para(doc, txt.STRUCTURAL_FE_INTRO, "Standard_Text")
        for t in txt.STRUCTURAL_FE_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STRUCTURAL_FE_SUB_BULLETS_1:
            _add_para(doc, t, "sub_bullets")
        for t in txt.STRUCTURAL_FE_SCOPE_2:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.STRUCTURAL_FE_SUB_BULLETS_2:
            _add_para(doc, t, "sub_bullets")
        for t in txt.STRUCTURAL_FE_SCOPE_3:
            _add_para(doc, t, "ASTUTE_Bullet_Points")

    if ca:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 5: Technical Support{_get_opt(ca)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        _add_para(doc, txt.CONSTRUCTION_ADVICE_INTRO, "Standard_Text")
        for t in txt.CONSTRUCTION_ADVICE_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.CONSTRUCTION_ADVICE_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        ml = ca.get("meetings_per_month")
        ml_text = "one meeting" if ml and str(ml) == "1" else f"{ml} meetings" if ml else "one meeting"
        _add_para(doc, f"Where necessary, attend fire engineering workshops or design team meetings to discuss the scheme (an average of {ml_text} per month is included in this scope); and", "ASTUTE_Bullet_Points")
        _add_para(doc, f"Where necessary, attend meetings with the approval authorities to seek approval for the detailed fire strategy (an average of {ml_text} per month is included in this scope). ", "ASTUTE_Bullet_Points")
        hrs = ca.get("hours_per_month")
        hrs_text = "one hour" if hrs and str(hrs) == "1" else f"{hrs} hours" if hrs else "N/A hours"
        _add_para(doc, f"The fee has been based on an average of {hrs_text} per month. Should more time than this be required we would reserve the right to consider this as a variation to the scope of services and request additional fees. ", "Standard_Text")

    if sv:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Site Visits{_get_opt(sv)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope", "sub_sub_heading")
        for t in txt.SITE_VISITS_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.SITE_VISITS_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")

    if po:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Phased Occupation Strategy{_get_opt(po)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope ", "sub_sub_heading")
        _add_para(doc, txt.PHASED_OCCUPATION_INTRO, "Standard_Text")
        for t in txt.PHASED_OCCUPATION_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.PHASED_OCCUPATION_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        _add_para(doc, txt.PHASED_OCCUPATION_DELIVERABLES_TEMPLATE.format(legislation=legislation), "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, po, "Where necessary, attend meetings with the project team and relevant stakeholders. ", "Attend all relevant meetings with the project team and relevant stakeholders")

    if cfsmp:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Construction Fire Safety Plan{_get_opt(cfsmp)}", "ASTUTE SubHeader")
        _add_para(doc, txt.CFSMP_INTRO, "Standard_Text")
        for t in txt.CFSMP_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.CFSMP_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        _add_meeting_text(doc, cfsmp, "Where necessary, attend meetings with the project team and relevant stakeholders. ", "Attend all relevant meetings with the project team and relevant stakeholders")

    if sr:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Construction Fire Risk Assessment{_get_opt(sr)}", "ASTUTE SubHeader")
        for t in txt.CONSTRUCTION_RA_TEXT:
            _add_para(doc, t, "Standard_Text")

    if r38:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Regulation 38 Information{_get_opt(r38)}", "ASTUTE SubHeader")
        for t in txt.REG38_TEXT:
            _add_para(doc, t, "Standard_Text")

    if ews1:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: EWS1 Forms{_get_opt(ews1)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope ", "sub_sub_heading")
        _add_para(doc, txt.EWS1_INTRO, "Standard_Text")
        _add_para(doc, "The works would comprise of: ", "Standard_Text")
        for t in txt.EWS1_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, txt.EWS1_FOLLOWING, "Standard_Text")
        for t in txt.EWS1_FOLLOWING_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Exclusions ", "sub_sub_heading")
        _add_para(doc, txt.EWS1_EXCLUSIONS_INTRO, "Standard_Text")
        for t in txt.EWS1_EXCLUSIONS:
            _add_para(doc, t, "ASTUTE_Bullet_Points")

    if rro:
        _add_para(doc, f"Scope of Works - RIBA Stage 5: Pre-Occupation Fire Risk Assessment{_get_opt(rro)}", "ASTUTE SubHeader")
        for t in txt.COMPLETION_RA_TEXT:
            _add_para(doc, t, "Standard_Text")

    if pr:
        _add_para(doc, f"Scope of Works \u2013 Third Party Review{_get_opt(pr)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope ", "sub_sub_heading")
        for t in txt.PEER_REVIEW_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        for t in txt.PEER_REVIEW_SUB_BULLETS:
            _add_para(doc, t, "sub_bullets")
        for t in txt.PEER_REVIEW_SCOPE_2:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_meeting_text(doc, pr, "Where necessary, attend meetings to discuss our comments on the fire engineering design. ", "Attendance at relevant meetings to discuss our comments on the fire engineering design.")
        _add_para(doc, "One revision of the peer review report should any changes to the study be made by the project fire engineer following discussion.", "ASTUTE_Bullet_Points")

    if cm:
        _add_para(doc, f"Scope of Works \u2013 RIBA Stage 5: Client Monitoring{_get_opt(cm)}", "ASTUTE SubHeader")
        _add_para(doc, "Fire Engineering Scope ", "sub_sub_heading")
        _add_para(doc, txt.CLIENT_MONITORING_INTRO, "Standard_Text")
        for t in txt.CLIENT_MONITORING_SCOPE:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, "Deliverables, Meeting Attendance and Completion", "sub_sub_heading")
        for t in txt.CLIENT_MONITORING_DELIVERABLES:
            _add_para(doc, t, "ASTUTE_Bullet_Points")
        _add_para(doc, txt.CLIENT_MONITORING_CLOSING, "Standard_Text")


def _write_appendix_b(doc, data, input_data):
    s14 = data.design_stages_1_4
    s6 = data.design_stages_6
    third_party = s14.peer_review.included

    if not third_party:
        _add_para(doc, txt.EXCLUSIONS_INTRO, "Standard_Text")

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("Structural Fire Engineering")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_STRUCTURAL_FE_INCLUDED if s14.warehouse_structural.included else txt.EXCL_STRUCTURAL_FE_NOT_INCLUDED)
        n2.font.underline = False

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("Evacuation Modelling")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_EVAC_MODELLING)
        n2.font.underline = False

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("Smoke Modelling")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_SMOKE_MODELLING_INCLUDED if (s14.open_plan_cfd.included or s14.common_corridor_cfd.included or s14.warehouse_cfd.included) else txt.EXCL_SMOKE_MODELLING_NOT_INCLUDED)
        n2.font.underline = False

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("Fire Strategy Drawings in CAD/BIM Format")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_FIRE_STRATEGY_DRAWINGS)
        n2.font.underline = False

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("External Walls")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_EXTERNAL_WALLS_WITH_EWS1 if s6.ews1_forms.included else txt.EXCL_EXTERNAL_WALLS_WITHOUT_EWS1)
        n2.font.underline = False

        if not (s14.stage_1.included or s14.stage_2.included or s14.stage_3.included or s14.stage_4.included):
            p = _add_para(doc, "", "ASTUTE_Bullet_Points")
            n = p.add_run("Fire Strategy Development")
            n.font.underline = True
            n2 = p.add_run(txt.EXCL_FIRE_STRATEGY_DEV)
            n2.font.underline = False

        p = _add_para(doc, "", "ASTUTE_Bullet_Points")
        n = p.add_run("Letters of Comfort")
        n.font.underline = True
        n2 = p.add_run(txt.EXCL_LETTERS_OF_COMFORT)
        n2.font.underline = False

        if _find_service(input_data, "site_visits"):
            p = _add_para(doc, "", "ASTUTE_Bullet_Points")
            n = p.add_run("Site Visit Records")
            n.font.underline = True
            n2 = p.add_run(txt.EXCL_SITE_VISIT_RECORDS)
            n2.font.underline = False

            p = _add_para(doc, "", "ASTUTE_Bullet_Points")
            n = p.add_run("Validity of Site Visit Observations")
            n.font.underline = True
            n2 = p.add_run(txt.EXCL_SITE_VISIT_VALIDITY)
            n2.font.underline = False
    else:
        _add_para(doc, txt.EXCLUSIONS_PEER_REVIEW, "Standard_Text")


def get_proposal_filename(data) -> str:
    today = date.today()
    date_str = today.strftime("%d_%m_%Y")
    project_name = data.project.project_name
    return f"{project_name} Fire Dynamics Fee Proposal {date_str}.docx"
