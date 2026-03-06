"""Generate EFS Word report using python-docx (plain document)."""

from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

from services.efs_calculator import run_single_calc


HEADERS = [
    "Elevation",
    "Boundary\nDist (m)",
    "ER Width\n(m)",
    "ER Height\n(m)",
    "BRE Width\n(m)",
    "BRE Height\n(m)",
    "BRE %\nUnprotected",
    "Allowable\nArea (m²)",
    "Actual\nArea (m²)",
    "Actual Allowable\nArea (m²)",
    "Actual %\nUnprotected",
]


def generate_efs_report(
    height_list: List[float],
    width_list: List[float],
    boundary_dist_list: List[float],
    has_suppression_list: List[bool],
    is_commercial: bool = True,
    building_fire_resistance_period: int = 60,
) -> BytesIO:
    """Generate an EFS Word report and return as BytesIO stream."""

    rows_data = []
    elevation_obj = []

    for idx in range(len(height_list)):
        height = height_list[idx]
        width = width_list[idx]
        boundary_dist = boundary_dist_list[idx]
        has_suppression = has_suppression_list[idx]

        percentage, actual_percentage, actual_area, allowable_area, actual_allowable_area, area, bre_height, bre_width = run_single_calc(
            height, width, boundary_dist, has_suppression, is_commercial
        )

        if boundary_dist < 1:
            bd_display = "<1"
            width_d = "-"
            height_d = "-"
            bre_height_d = "-"
            bre_width_d = "-"
            percentage = 0
            percentage_s = "0%"
            allowable_area_d = "-"
            actual_area_d = "-"
            actual_allowable_area_d = "-"
            actual_percentage_s = "-"
        elif boundary_dist == 1:
            bd_display = str(boundary_dist)
            width_d = "-"
            height_d = "-"
            bre_height_d = "-"
            bre_width_d = "-"
            percentage_s = "0%"
            allowable_area_d = "-"
            actual_area_d = "-"
            actual_allowable_area_d = "-"
            actual_percentage_s = "-"
        elif percentage == 100:
            bd_display = str(boundary_dist)
            width_d = str(width)
            height_d = str(height)
            bre_height_d = str(bre_height)
            bre_width_d = str(bre_width)
            percentage_s = "100%"
            allowable_area_d = "-"
            actual_area_d = "-"
            actual_allowable_area_d = "-"
            actual_percentage_s = "-"
        else:
            bd_display = ">120" if boundary_dist > 120 else str(boundary_dist)
            width_d = str(width)
            height_d = str(height)
            bre_height_d = str(bre_height)
            bre_width_d = str(bre_width)
            percentage_s = str(round(percentage, 1)) + "%"
            allowable_area_d = str(round(allowable_area, 1))
            actual_area_d = str(actual_area)
            actual_allowable_area_d = str(actual_allowable_area)
            actual_percentage_s = str(actual_percentage) + "%"

        rows_data.append([
            str(idx + 1), bd_display, width_d, height_d,
            bre_width_d, bre_height_d, percentage_s,
            allowable_area_d, actual_area_d,
            actual_allowable_area_d, actual_percentage_s,
        ])

        elevation_obj.append({
            "BRE_perc": percentage_s,
            "perc": actual_percentage_s,
            "b_dist": bd_display,
            "less_1m": boundary_dist < 1,
        })

    # Build conditional text
    unprotected_bool_list = [
        obj["BRE_perc"] == "100%" or obj["perc"] == "100.0%"
        for obj in elevation_obj
    ]

    elevations_less_1m_bool_list = [obj["less_1m"] for obj in elevation_obj]
    elevations_less_1m_list = [
        str(i + 1) for i, v in enumerate(elevations_less_1m_bool_list) if v
    ]

    has_elevations_protected_bool_list = [
        not unprotected_bool_list[i] and not elevations_less_1m_bool_list[i]
        for i in range(len(unprotected_bool_list))
    ]
    elevations_protected_list = [
        str(i + 1) for i, v in enumerate(has_elevations_protected_bool_list) if v
    ]
    elevations_protected_percentage_list = [
        elevation_obj[i]["perc"]
        for i, v in enumerate(has_elevations_protected_bool_list) if v
    ]

    some_unprotected_list = [
        str(i + 1) for i, v in enumerate(unprotected_bool_list) if v
    ]

    frp = building_fire_resistance_period

    # --- Build document ---
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('External Fire Spread Assessment', level=1)
    doc.add_paragraph(
        f'Building type: {"Commercial" if is_commercial else "Residential"}'
    )

    # Results table
    doc.add_heading('Results', level=2)

    table = doc.add_table(rows=1 + len(rows_data), cols=len(HEADERS))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Narrative text
    doc.add_heading('Commentary', level=2)

    if elevations_protected_list:
        if len(elevations_protected_list) == 1:
            plural = "elevation"
            this_text = "this elevation"
            list_text = elevations_protected_list[0]
            perc_text = str(elevations_protected_percentage_list[0])
        else:
            plural = "elevations"
            this_text = "these elevations"
            list_text = ", ".join(elevations_protected_list[:-1]) + " and " + elevations_protected_list[-1]
            perc_text = ", ".join(str(p) for p in elevations_protected_percentage_list[:-1]) + " and " + str(elevations_protected_percentage_list[-1])
        doc.add_paragraph(
            f"The results of the external fire spread assessment show that {plural} {list_text} "
            f"should be designed to achieve {perc_text} fire resistance to prevent external fire spread "
            f"across the site boundary. Specifically, {this_text} should be designed to achieve a fire "
            f"resistance period of {frp} minutes integrity and 15 minutes insulation. The remainder of "
            f"these elevations are not required to be designed to achieve fire resistance."
        )

    if elevations_less_1m_list:
        el_plural = "elevation" if len(elevations_less_1m_list) == 1 else "elevations"
        el_list_text = (
            elevations_less_1m_list[0]
            if len(elevations_less_1m_list) == 1
            else ", ".join(elevations_less_1m_list[:-1]) + " and " + elevations_less_1m_list[-1]
        )
        doc.add_paragraph(
            f"Given that the distance to the relevant boundary on {el_plural} {el_list_text} is less than 1m, "
            f"the whole elevation should be designed to achieve a fire resistance period of {frp} minutes "
            f"integrity and {frp} minutes insulation. It is permitted to have some small areas of the elevation "
            f"which do not meet this requirement as long as their size and spacing is as per Diagram 13.5 of "
            f"AD-B (extracted below)."
        )

    if some_unprotected_list:
        s_plural = "elevation" if len(some_unprotected_list) == 1 else "elevations"
        s_list_text = (
            some_unprotected_list[0]
            if len(some_unprotected_list) == 1
            else ", ".join(some_unprotected_list[:-1]) + " and " + some_unprotected_list[-1]
        )
        doc.add_paragraph(
            f"The table shows that there is no requirement for {s_plural} {s_list_text} to be designed to "
            f"achieve a particular period of fire resistance to prevent external fire spread across the site boundary."
        )

    if not elevations_protected_list and not elevations_less_1m_list and not some_unprotected_list:
        doc.add_paragraph("No external fire spread requirements identified.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
