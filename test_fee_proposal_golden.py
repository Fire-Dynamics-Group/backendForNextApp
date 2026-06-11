"""Golden-file regression baseline for fee proposal generation.

This is a CHARACTERIZATION test: it locks the current rendered output of
``generate_proposal`` so that the upcoming refactor (moving narrative text from
Python constants into the database) can be proven to produce byte-identical
documents. It is not red-green TDD - it is a safety net for the cutover.

The generated ``.docx`` embeds ``date.today()`` / ``datetime.now()``, so dates
are frozen via monkeypatch to keep the snapshots deterministic.

Baselines live as JSON under ``golden/``. To (re)generate them after an
intentional change, run with ``UPDATE_GOLDEN=1``.
"""
import io
import json
import os
from datetime import date, datetime

import pytest
from docx import Document

import services.fee_document_service as svc
from services.fee_document_service import generate_proposal, get_proposal_filename
from models.fee_proposal_models import (
    FeeProposalRequest,
    ClientDetails,
    ProjectDetails,
    FeeOptions,
    DesignStagesRiba1to4,
    DesignStagesRiba5,
    DesignStagesRiba6,
    ServiceConfig,
    CountryEnum,
)

GOLDEN_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "golden")
FROZEN_DATE = date(2026, 6, 11)
FROZEN_DATETIME = datetime(2026, 6, 11, 12, 0, 0)


class _FrozenDate(date):
    @classmethod
    def today(cls):
        return FROZEN_DATE


class _FrozenDateTime(datetime):
    @classmethod
    def now(cls, tz=None):
        return FROZEN_DATETIME


@pytest.fixture(autouse=True)
def _freeze_time(monkeypatch):
    """Freeze date.today() and datetime.now() inside the document service."""
    monkeypatch.setattr(svc, "date", _FrozenDate)
    monkeypatch.setattr(svc, "datetime", _FrozenDateTime)


def _run_repr(run):
    return {
        "text": run.text,
        "bold": run.bold,
        "underline": run.underline,
        "superscript": bool(run.font.superscript),
        "image": "graphicData" in run._element.xml,
    }


def _para_repr(p):
    return {
        "style": p.style.name if p.style is not None else None,
        "text": p.text,
        "runs": [_run_repr(r) for r in p.runs],
    }


def _extract(buf: io.BytesIO) -> dict:
    """Canonical, comparable representation of a generated proposal."""
    buf.seek(0)
    doc = Document(buf)
    paragraphs = [_para_repr(p) for p in doc.paragraphs]
    tables = []
    for table in doc.tables:
        tables.append({
            "style": table.style.name if table.style is not None else None,
            "rows": [
                [
                    {"text": cell.text,
                     "style": cell.paragraphs[0].style.name if cell.paragraphs and cell.paragraphs[0].style else None}
                    for cell in row.cells
                ]
                for row in table.rows
            ],
        })
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    return {"header": header_text, "paragraphs": paragraphs, "tables": tables}


def _check_golden(name: str, data: dict):
    os.makedirs(GOLDEN_DIR, exist_ok=True)
    path = os.path.join(GOLDEN_DIR, f"{name}.json")
    serialized = json.dumps(data, indent=2, ensure_ascii=False)
    if os.environ.get("UPDATE_GOLDEN") or not os.path.exists(path):
        with open(path, "w", encoding="utf-8") as f:
            f.write(serialized)
        pytest.skip(f"Wrote golden baseline {name}.json (no prior baseline to compare).")
    with open(path, "r", encoding="utf-8") as f:
        expected = json.load(f)
    assert data == expected, f"Output diverged from golden baseline {name}.json"


def _request(**kwargs) -> FeeProposalRequest:
    base = dict(
        client=ClientDetails(first_name="Test", surname="Client",
                             address_lines=["1 Test Street", "Testville", "TE1 1ST"]),
        project=ProjectDetails(project_name="Test Tower", project_location="London",
                               country=CountryEnum.ENGLAND_WALES),
        fee_options=FeeOptions(engineer_name="Sam Bennett", pii_limit=100000,
                               include_hourly_rates=False),
    )
    base.update(kwargs)
    return FeeProposalRequest(**base)


# --- Tracer fixture: kitchen sink (every non-peer-review service included) ---

def _kitchen_sink() -> FeeProposalRequest:
    return _request(
        fee_options=FeeOptions(engineer_name="Sam Bennett", pii_limit=250000,
                               include_hourly_rates=True),
        design_stages_1_4=DesignStagesRiba1to4(
            stage_1=ServiceConfig(included=True, fee=5000, limit_meetings=True, meeting_number=3),
            stage_2=ServiceConfig(included=True, fee=7500, end_date_month="June", end_date_year="2026"),
            london_plan=ServiceConfig(included=True, fee=3000),
            gateway=ServiceConfig(included=True, fee=3500),
            stage_3=ServiceConfig(included=True, fee=9000),
            stage_4=ServiceConfig(included=True, fee=11000),
            common_corridor_cfd=ServiceConfig(included=True, fee=8000, num_models=3),
            open_plan_cfd=ServiceConfig(included=True, fee=8500, num_models=2),
            warehouse_cfd=ServiceConfig(included=True, fee=12000, num_models=4),
            warehouse_structural=ServiceConfig(included=True, fee=14000),
        ),
        design_stages_5=DesignStagesRiba5(
            construction_advice=ServiceConfig(included=True, fee=6000, hours_per_month=5, meetings_per_month=1),
            site_visits=ServiceConfig(included=True, fee=2000),
            site_risk_assessment=ServiceConfig(included=True, fee=2500),
            cfsmp=ServiceConfig(included=True, fee=3000),
            phased_occupation=ServiceConfig(included=True, fee=3500),
            client_monitoring=ServiceConfig(included=True, fee=4000),
        ),
        design_stages_6=DesignStagesRiba6(
            regulation_38=ServiceConfig(included=True, fee=1500),
            ews1_forms=ServiceConfig(included=True, fee=1800),
            rro_risk_assessment=ServiceConfig(included=True, fee=2200),
        ),
    )


def test_golden_kitchen_sink():
    data = _extract(generate_proposal(_kitchen_sink()))
    data["filename"] = get_proposal_filename(_kitchen_sink())
    _check_golden("kitchen_sink", data)
