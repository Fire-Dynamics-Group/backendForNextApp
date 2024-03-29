# base file -> use a template word doce
# receive radiation data from frontend
# send back appendix word file
# LATER: image to be sent from frontend & inserted into word doc
import io
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from constants import font_name_light



document_name = "Oil Pan Fire Appendix - Template.docx"
document_path = Path(__file__).parent /"Word Templates"/document_name
doc = DocxTemplate(document_path)

def Delete_row_in_table(table, row, document):
    if type(table) == int:
        table_object = document.tables[table]
    else:
        table_object = table
    table_object._tbl.remove(table_object.rows[row]._tr)

def replace_table_cell_content(cell, replacement_text, is_bold=False, alignment=1):
    # later make object with fonts etc
    cell.text = replacement_text
    paragraphs = cell.paragraphs
    paragraphs[0].alignment = alignment # 1 = centered
    # for run in paragraphs.runs:
    run = paragraphs[0].runs
    font = run[0].font
    font.size = Pt(9) # pull from object
    font.name = font_name_light
    font.color.rgb = RGBColor(64,64,64) #gray color # to be pulled from object
    font.bold = is_bold
# needs totalHeatFlux=472, radiantHeatEndpoint = 1.3333
def fillWordDoc(
                timeArray, 
                accumulatedDistanceList, 
                hobDistanceList, 
                qList, 
                timestepFEDList, 
                accumulatedFEDList,
                # TODO: SEND FROM FRONTEND
                totalHeatFlux=476,
                walkingSpeed=1.2,
                doorOpeningDuration=None,
                output_filename="Oil Pan Fire Appendix.docx"           
        ):
    #     # CHIP_PAN_ALLOWED=True, HAS_CUSTOM_FIRE_SIZE=False
    # fireSizeObject = [
    #   {"description": "Chip Pan", "size": 476},
    #   {"description": "Chip Pan Banned", "size": 150.5},
    # ] 
    print(output_filename, "rad.py")
    if totalHeatFlux == 476:
        CHIP_PAN_ALLOWED = True
        HAS_CUSTOM_FIRE_SIZE = False
    elif totalHeatFlux == 150.5:
        CHIP_PAN_ALLOWED = False
        HAS_CUSTOM_FIRE_SIZE = False
    else:
        CHIP_PAN_ALLOWED = False
        HAS_CUSTOM_FIRE_SIZE = True

    HAS_DOOR = True
    if doorOpeningDuration == 11:
        HAS_CUSTOM_DOOR_DURATION = False
    elif doorOpeningDuration == None:
        HAS_CUSTOM_DOOR_DURATION = False
        HAS_DOOR = False        
    else:
        HAS_CUSTOM_DOOR_DURATION = True

    
    context = {
        "CHIP_PAN_ALLOWED": CHIP_PAN_ALLOWED,
        "HAS_CUSTOM_FIRE_SIZE": HAS_CUSTOM_FIRE_SIZE,
        "FIRE_Q": totalHeatFlux,
        "THIRD_FIRE_Q": round(totalHeatFlux/3, 1),
        "HAS_CUSTOM_WALKING_SPEED": False,
        "WALKING_SPEED": walkingSpeed,
        "HAS_CUSTOM_DOOR_DURATION": HAS_CUSTOM_DOOR_DURATION,
        "DOOR_DURATION": doorOpeningDuration,
        "HAS_DOOR": HAS_DOOR,
        "MAX_FED": round(max(accumulatedFEDList), 3)
    }

    doc.render(context)

    doc.save(output_filename)
    
    document = Document(output_filename)


    def alter_table_rows(total_rows, table, document, header_rows = 1, has_door=False):
        # header should be increase by 1
        current_rows = len(table.rows)
        if has_door:
            header_rows += 1
        required_rows = total_rows + header_rows
        # remove bottom rows
        rows_to_remove = current_rows - required_rows
        # loop remove row[-1] * rows_to_remove
        if not has_door:
            row = -1
        else: # leave last two rows intact
            row = -4
            rows_to_remove -= 2
        for i in range(rows_to_remove):
            # TODO: if has door -> last two rows to remain
            # 
            if not has_door and i == rows_to_remove - 1:
                row = 1
            Delete_row_in_table(table, row=row, document=document)    

    def fill_radiation_table(table_object, has_door=False):
        # LATER include final line for door opening with n/a for distance travelled
        # 
        table_rows = table_object.rows
        if has_door:
            start_row = 2
            end = len(table_rows) 
            # should also skip -3 index, fill -2 and then repeat totalFED for final col -1
        else:
            start_row = 1
            end = len(table_rows)
        list_index = 0

        # perhaps list of indexes needed
        # repeat cumulative for total row when door present
        for row_index in range(start_row,end):
            if has_door and row_index == len(table_rows) - 3:
                pass
            else:
                # list_index = row_index - 1
                for cell_idx,target_cell in enumerate(table_rows[row_index].cells):
                    if has_door and row_index == len(table_rows) - 1:
                        if cell_idx==len(table_rows[row_index].cells) - 1:
                            cell_text = round(accumulatedFEDList[-1], 2)
                            replace_table_cell_content(cell=target_cell, replacement_text=str((cell_text)))
                    else:
                        if cell_idx == 0:
                            cell_text = timeArray[list_index]
                        elif cell_idx == 1:
                            cell_text = round(accumulatedDistanceList[list_index], 2)
                        elif cell_idx == 2:
                            cell_text = round(hobDistanceList[list_index], 2)
                        elif cell_idx == 3:
                            cell_text = round(qList[list_index], 2)
                        elif cell_idx == 4:
                            cell_text = round(timestepFEDList[list_index], 2)
                        elif cell_idx == 5:
                            cell_text = round(accumulatedFEDList[list_index], 2)

                        replace_table_cell_content(cell=target_cell, replacement_text=str((cell_text)))
                
                list_index += 1
                if list_index == len(timeArray):
                    pass 
                
    radiation_table = document.tables[1] # later find tables by name below??
    alter_table_rows(total_rows=len(timeArray), document=document, table=radiation_table, has_door=HAS_DOOR)
    # function to input all time, distance,...etc
    fill_radiation_table(table_object=radiation_table, has_door=HAS_DOOR)

    # rad_door_table = document.tables[2]
    # alter_table_rows(total_rows=len(timeArray), document=document, table=rad_door_table, has_door=True)
    # fill_radiation_table(rad_door_table, has_door=True)


    document.save(output_filename)

    bytes_io = io.BytesIO()
    document.save(bytes_io)

    bytes_io.seek(0)    
    return bytes_io
if __name__ == '__main__':

    bytes_io = fillWordDoc(
        timeArray=[        
            0,
            1,
            2,
            3,
            4,
            5,
            6,
            7,
            8,
            9,
            10,
            10.35
            ],
        accumulatedDistanceList=[
            0,
            1.2,
            2.4,
            3.5999999999999996,
            5.120459156782531,
            6.3204591567825315,
            7.520459156782532,
            9.352439756686653,
            10.552439756686653,
            11.752439756686652,
            12.967707134626034,
            13.798241890504794        
        ],
        hobDistanceList=[
            6.363175307973213,
            5.175998158550895,
            3.9964876922923867,
            2.8342319615668,
            2.3338112370936126,
            2.577102339521506,
            3.2735665636045987,
            4.1480441286038765,
            5.104823209281279,
            6.147533643924898,
            6.0038604284247725,
            6.003332407921453        
        ],
        qList=[
            0.31180611318206564,
            0.47124220132480643,
            0.7904518908522363,
            1.5716709278996468,
            2.3179333771197017,
            1.9009429005330298,
            0,
            0,
            0,
            0,
            0,
            0        
        ],
        timestepFEDList=[
            0.0026533065097377223,
            0.0045955302020744765,
            0.009143136621402696,
            0.022807717489928544,
            0.03823887730483287,
            0.02937309677957744,
            0,
            0,
            0,
            0,
            0,
            0        
        ],
        accumulatedFEDList=[
            0.0026533065097377223,
            0.007248836711812199,
            0.016391973333214897,
            0.03919969082314344,
            0.07743856812797631,
            0.10681166490755375,
            0.10681166490755375,
            0.10681166490755375,
            0.10681166490755375,
            0.10681166490755375,
            0.10681166490755375,
            0.10681166490755375        
        ]
    )

