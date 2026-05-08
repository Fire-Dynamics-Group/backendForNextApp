"""Tests for fee_document_service - verifies expected end date is rendered
in both single-service and multi-service fee proposals."""
import io
import pytest
from docx import Document

from models.fee_proposal_models import (
    FeeProposalRequest,
    ClientDetails,
    ProjectDetails,
    FeeOptions,
    DesignStagesRiba1to4,
    ServiceConfig,
    CountryEnum,
)
from services.fee_document_service import generate_proposal


def _base_request(stages_1_4: DesignStagesRiba1to4) -> FeeProposalRequest:
    return FeeProposalRequest(
        client=ClientDetails(first_name="Test", surname="Client", address_lines=["1 Test St"]),
        project=ProjectDetails(project_name="Test Project", project_location="London", country=CountryEnum.ENGLAND_WALES),
        fee_options=FeeOptions(engineer_name="Sam Bennett", pii_limit=100000, include_hourly_rates=False),
        design_stages_1_4=stages_1_4,
    )


def _all_text(buf: io.BytesIO) -> str:
    buf.seek(0)
    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


class TestExpectedEndDateRendering:
    def test_single_service_with_end_date_renders_expected_end_date(self):
        stages = DesignStagesRiba1to4(
            stage_1=ServiceConfig(included=True, fee=5000, end_date_month="March", end_date_year="2026"),
        )
        buf = generate_proposal(_base_request(stages))
        text = _all_text(buf)
        assert "(expected end date March 2026)" in text

    def test_single_service_with_end_date_appears_after_exc_vat(self):
        stages = DesignStagesRiba1to4(
            stage_1=ServiceConfig(included=True, fee=5000, end_date_month="March", end_date_year="2026"),
        )
        buf = generate_proposal(_base_request(stages))
        text = _all_text(buf)
        assert "exc. VAT (expected end date March 2026)." in text

    def test_single_service_without_end_date_unchanged(self):
        stages = DesignStagesRiba1to4(stage_1=ServiceConfig(included=True, fee=5000))
        buf = generate_proposal(_base_request(stages))
        text = _all_text(buf)
        assert "expected end date" not in text
        assert "exc. VAT." in text

    def test_multi_service_renders_expected_end_date_per_row(self):
        stages = DesignStagesRiba1to4(
            stage_1=ServiceConfig(included=True, fee=5000, end_date_month="March", end_date_year="2026"),
            stage_2=ServiceConfig(included=True, fee=7500, end_date_month="June", end_date_year="2026"),
        )
        buf = generate_proposal(_base_request(stages))
        text = _all_text(buf)
        assert "RIBA Stage 1 (expected end date March 2026)" in text
        assert "RIBA Stage 2 (expected end date June 2026)" in text

    def test_legacy_up_to_phrasing_is_gone(self):
        stages = DesignStagesRiba1to4(
            stage_1=ServiceConfig(included=True, fee=5000, end_date_month="March", end_date_year="2026"),
            stage_2=ServiceConfig(included=True, fee=7500, end_date_month="June", end_date_year="2026"),
        )
        buf = generate_proposal(_base_request(stages))
        text = _all_text(buf)
        assert "(Up to" not in text
